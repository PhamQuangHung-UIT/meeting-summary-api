from typing import Dict, Any
from datetime import datetime
import json
import io
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import zipfile

from app.utils.database import supabase
from app.services.recording_service import RecordingService
from app.services.transcript_service import TranscriptService
from app.services.summary_service import SummaryService


class ExportProcessor:
    def __init__(self, job: Dict[str, Any]):
        self.job = job
        self.recording_id = job['recording_id']
        self.export_type = job['export_type']
        self.export_id = job['export_id']

    def process(self) -> str:
        """Process export and return file_path in storage"""

        if self.export_type == "TRANSCRIPT_PDF":
            return self._export_transcript_pdf()
        elif self.export_type == "TRANSCRIPT_DOCX":
            return self._export_transcript_docx()
        elif self.export_type == "SUMMARY_PDF":
            return self._export_summary_pdf()
        elif self.export_type == "SUMMARY_DOCX":
            return self._export_summary_docx()
        elif self.export_type == "FULL_ZIP":
            return self._export_full_zip()
        else:
            raise ValueError(f"Unsupported export type: {self.export_type}")

    def _get_transcript_data(self):
        """Get transcript with segments"""
        transcripts = TranscriptService.get_transcripts_by_recording_id(
            self.recording_id,
            latest=True
        )
        if not transcripts:
            raise ValueError("No transcript found")

        transcript = TranscriptService.get_transcript_by_id(transcripts[0]['transcript_id'])
        return transcript

    def _get_summary_data(self):
        """Get latest summary"""
        summaries = SummaryService.get_summaries_by_recording_id(
            self.recording_id,
            latest=True
        )
        if not summaries:
            raise ValueError("No summary found")

        return summaries[0]

    def _get_recording_data(self):
        """Get recording details"""
        recording = RecordingService.get_recording_by_id(self.recording_id)
        if not recording:
            raise ValueError("Recording not found")
        return recording

    def _export_transcript_pdf(self) -> str:
        """Generate PDF from transcript"""
        transcript = self._get_transcript_data()
        recording = self._get_recording_data()

        # Create PDF in memory
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
                                rightMargin=72, leftMargin=72,
                                topMargin=72, bottomMargin=18)

        # Container for the 'Flowable' objects
        elements = []

        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='#1a1a1a',
            spaceAfter=30,
            alignment=TA_CENTER
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor='#333333',
            spaceAfter=12
        )

        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            leading=16,
            spaceAfter=10
        )

        # Add title
        elements.append(Paragraph(f"Transcript: {recording['title']}", title_style))
        elements.append(Spacer(1, 0.2 * inch))

        # Add metadata
        created_date = datetime.fromisoformat(recording['created_at'].replace('Z', '+00:00'))
        elements.append(Paragraph(f"<b>Created:</b> {created_date.strftime('%Y-%m-%d %H:%M')}", normal_style))
        elements.append(Paragraph(f"<b>Duration:</b> {recording['duration_seconds']:.2f} seconds", normal_style))
        elements.append(Paragraph(f"<b>Version:</b> {transcript.version_no}", normal_style))
        elements.append(Spacer(1, 0.3 * inch))

        # Add segments
        elements.append(Paragraph("Transcript Content", heading_style))
        elements.append(Spacer(1, 0.1 * inch))

        for segment in transcript.segments:
            time_str = f"[{self._format_time(segment['start_time'])} - {self._format_time(segment['end_time'])}]"
            speaker = segment.get('speaker_label', 'Unknown')
            content = segment['content']

            text = f"<b>{speaker}</b> {time_str}<br/>{content}"
            elements.append(Paragraph(text, normal_style))
            elements.append(Spacer(1, 0.1 * inch))

        # Build PDF
        doc.build(elements)

        # Upload to storage
        buffer.seek(0)
        filename = f"{self.export_id}_transcript.pdf"
        file_path = f"{self.recording_id}/{filename}"

        supabase.storage.from_("exports").upload(
            file_path,
            buffer.getvalue(),
            file_options={"content-type": "application/pdf"}
        )

        return file_path

    def _export_transcript_docx(self) -> str:
        """Generate DOCX from transcript"""
        transcript = self._get_transcript_data()
        recording = self._get_recording_data()

        # Create document
        doc = Document()

        # Add title
        title = doc.add_heading(f"Transcript: {recording['title']}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata
        doc.add_paragraph(f"Created: {recording['created_at']}")
        doc.add_paragraph(f"Duration: {recording['duration_seconds']:.2f} seconds")
        doc.add_paragraph(f"Version: {transcript.version_no}")
        doc.add_paragraph()

        # Add segments
        doc.add_heading("Transcript Content", level=1)

        for segment in transcript.segments:
            time_str = f"[{self._format_time(segment['start_time'])} - {self._format_time(segment['end_time'])}]"
            speaker = segment.get('speaker_label', 'Unknown')

            p = doc.add_paragraph()
            p.add_run(f"{speaker} {time_str}\n").bold = True
            p.add_run(segment['content'])
            p.add_run("\n")

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Upload to storage
        filename = f"{self.export_id}_transcript.docx"
        file_path = f"{self.recording_id}/{filename}"

        supabase.storage.from_("exports").upload(
            file_path,
            buffer.getvalue(),
            file_options={"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
        )

        return file_path

    def _export_summary_pdf(self) -> str:
        """Generate PDF from summary"""
        summary = self._get_summary_data()
        recording = self._get_recording_data()

        # Create PDF
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
                                rightMargin=72, leftMargin=72,
                                topMargin=72, bottomMargin=18)

        elements = []
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='#1a1a1a',
            spaceAfter=30,
            alignment=TA_CENTER
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor='#333333',
            spaceAfter=12
        )

        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            leading=16,
            spaceAfter=10
        )

        # Add title
        elements.append(Paragraph(f"Summary: {recording['title']}", title_style))
        elements.append(Spacer(1, 0.2 * inch))

        # Add metadata
        elements.append(Paragraph(f"<b>Created:</b> {summary['created_at']}", normal_style))
        elements.append(Paragraph(f"<b>Style:</b> {summary.get('summary_style', 'N/A')}", normal_style))
        elements.append(Spacer(1, 0.3 * inch))

        # Parse and add summary content
        content = summary.get('content_structure', {})

        # Overview
        if 'overview' in content:
            elements.append(Paragraph("Overview", heading_style))
            elements.append(Paragraph(content['overview'], normal_style))
            elements.append(Spacer(1, 0.2 * inch))

        # Key Points
        if 'key_points' in content:
            elements.append(Paragraph("Key Points", heading_style))
            for point in content['key_points']:
                elements.append(Paragraph(f"• {point}", normal_style))
            elements.append(Spacer(1, 0.2 * inch))

        # Action Items
        if 'action_items' in content:
            elements.append(Paragraph("Action Items", heading_style))
            for item in content['action_items']:
                elements.append(Paragraph(f"• {item}", normal_style))

        doc.build(elements)

        # Upload to storage
        buffer.seek(0)
        filename = f"{self.export_id}_summary.pdf"
        file_path = f"{self.recording_id}/{filename}"

        supabase.storage.from_("exports").upload(
            file_path,
            buffer.getvalue(),
            file_options={"content-type": "application/pdf"}
        )

        return file_path

    def _export_summary_docx(self) -> str:
        """Generate DOCX from summary"""
        summary = self._get_summary_data()
        recording = self._get_recording_data()

        doc = Document()

        # Add title
        title = doc.add_heading(f"Summary: {recording['title']}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata
        doc.add_paragraph(f"Created: {summary['created_at']}")
        doc.add_paragraph(f"Style: {summary.get('summary_style', 'N/A')}")
        doc.add_paragraph()

        # Parse and add content
        content = summary.get('content_structure', {})

        if 'overview' in content:
            doc.add_heading("Overview", level=1)
            doc.add_paragraph(content['overview'])

        if 'key_points' in content:
            doc.add_heading("Key Points", level=1)
            for point in content['key_points']:
                doc.add_paragraph(f"• {point}")

        if 'action_items' in content:
            doc.add_heading("Action Items", level=1)
            for item in content['action_items']:
                doc.add_paragraph(f"• {item}")

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Upload
        filename = f"{self.export_id}_summary.docx"
        file_path = f"{self.recording_id}/{filename}"

        supabase.storage.from_("exports").upload(
            file_path,
            buffer.getvalue(),
            file_options={"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
        )

        return file_path

    def _export_full_zip(self) -> str:
        """Generate ZIP with transcript and summary"""
        # Create ZIP in memory
        buffer = io.BytesIO()

        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
            # Add transcript PDF
            try:
                transcript_path = self._export_transcript_pdf()
                transcript_data = supabase.storage.from_("exports").download(transcript_path)
                zipf.writestr("transcript.pdf", transcript_data)
            except Exception as e:
                print(f"Error adding transcript to ZIP: {e}")

            # Add summary PDF
            try:
                summary_path = self._export_summary_pdf()
                summary_data = supabase.storage.from_("exports").download(summary_path)
                zipf.writestr("summary.pdf", summary_data)
            except Exception as e:
                print(f"Error adding summary to ZIP: {e}")

        buffer.seek(0)

        # Upload ZIP
        filename = f"{self.export_id}_full.zip"
        file_path = f"{self.recording_id}/{filename}"

        supabase.storage.from_("exports").upload(
            file_path,
            buffer.getvalue(),
            file_options={"content-type": "application/zip"}
        )

        return file_path

    @staticmethod
    def _format_time(seconds: float) -> str:
        """Format seconds to MM:SS"""
        minutes = int(seconds // 60)
        secs = int(seconds % 60)
        return f"{minutes:02d}:{secs:02d}"